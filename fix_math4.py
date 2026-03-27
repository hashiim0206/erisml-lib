#!/usr/bin/env python3
"""
Fix Type 4: Display equations (m:oMathPara) that are entirely missing
from the HTML. These were standalone equation paragraphs in the docx
that got completely dropped during conversion.

Also fixes remaining inline single-space gaps.

Strategy: Walk through the docx paragraph by paragraph. For each
display-math paragraph, find the surrounding text paragraphs in
both docx and HTML to locate where the equation should be inserted.
"""

import sys, io, re, os, html as html_mod
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
import lxml.etree as ET

BOOK_DIR = r'C:\source\erisml-lib\docs\book'
DOCX_PATH = r'C:\source\erisml-lib\docs\papers\foundations\Geometric Ethics - The Mathematical Structure of Moral Reasoning - Bond - v1.14 - Mar 2026.docx'

ns = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
      'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def omath_to_unicode_inner(elem):
    parts = []
    def process(el):
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 't': parts.append(el.text or '')
        elif tag in ('r', 'e', 'sub', 'sup', 'num', 'den', 'deg'):
            for child in el: process(child)
        elif tag.endswith('Pr'): pass
        else:
            for child in el: process(child)
    process(elem)
    return ''.join(parts).strip()


def omath_to_html(elem):
    parts = []
    mns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    def process(el):
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 't': parts.append(html_mod.escape(el.text or ''))
        elif tag == 'sSub':
            base = el.find('m:e', ns); sub = el.find('m:sub', ns)
            if base is not None: process(base)
            if sub is not None:
                parts.append(f'<sub>{html_mod.escape(omath_to_unicode_inner(sub))}</sub>')
        elif tag == 'sSup':
            base = el.find('m:e', ns); sup = el.find('m:sup', ns)
            if base is not None: process(base)
            if sup is not None:
                parts.append(f'<sup>{html_mod.escape(omath_to_unicode_inner(sup))}</sup>')
        elif tag == 'sSubSup':
            base = el.find('m:e', ns); sub = el.find('m:sub', ns); sup = el.find('m:sup', ns)
            if base is not None: process(base)
            if sub is not None:
                parts.append(f'<sub>{html_mod.escape(omath_to_unicode_inner(sub))}</sub>')
            if sup is not None:
                parts.append(f'<sup>{html_mod.escape(omath_to_unicode_inner(sup))}</sup>')
        elif tag == 'f':
            num = el.find('m:num', ns); den = el.find('m:den', ns)
            n = html_mod.escape(omath_to_unicode_inner(num)) if num is not None else ''
            d = html_mod.escape(omath_to_unicode_inner(den)) if den is not None else ''
            parts.append(f'({n})/({d})')
        elif tag == 'rad':
            base = el.find('m:e', ns)
            parts.append('\u221a(')
            if base is not None: process(base)
            parts.append(')')
        elif tag == 'nary':
            chr_el = el.find('m:naryPr/m:chr', ns)
            op = chr_el.get(f'{{{mns}}}val', '\u2211') if chr_el is not None else '\u2211'
            sub = el.find('m:sub', ns); sup = el.find('m:sup', ns); base = el.find('m:e', ns)
            parts.append(op)
            if sub is not None:
                st = omath_to_unicode_inner(sub)
                if st.strip(): parts.append(f'<sub>{html_mod.escape(st)}</sub>')
            if sup is not None:
                st = omath_to_unicode_inner(sup)
                if st.strip(): parts.append(f'<sup>{html_mod.escape(st)}</sup>')
            parts.append(' ')
            if base is not None: process(base)
        elif tag == 'd':
            dPr = el.find('m:dPr', ns); beg, end = '(', ')'
            if dPr is not None:
                b = dPr.find('m:begChr', ns); e = dPr.find('m:endChr', ns)
                if b is not None: beg = b.get(f'{{{mns}}}val', '(')
                if e is not None: end = e.get(f'{{{mns}}}val', ')')
            parts.append(html_mod.escape(beg))
            e_els = el.findall('m:e', ns)
            for idx, ee in enumerate(e_els):
                if idx > 0: parts.append(', ')
                process(ee)
            parts.append(html_mod.escape(end))
        elif tag == 'acc':
            base = el.find('m:e', ns)
            if base is not None: process(base)
        elif tag == 'bar':
            base = el.find('m:e', ns)
            if base is not None: process(base)
        elif tag == 'eqArr':
            for idx, ee in enumerate(el.findall('m:e', ns)):
                if idx > 0: parts.append('<br>')
                process(ee)
        elif tag in ('r', 'e', 'sub', 'sup', 'num', 'den', 'deg', 'oMath', 'oMathPara'):
            for child in el: process(child)
        elif tag.endswith('Pr'): pass
        else:
            for child in el: process(child)
    process(elem)
    return ''.join(parts).strip()


def normalize(text):
    t = re.sub(r'\s+', ' ', text).strip().lower()
    t = t.replace('\u2013', '-').replace('\u2014', '-')
    t = t.replace('\u2019', "'").replace('\u2018', "'")
    t = t.replace('\u201c', '"').replace('\u201d', '"')
    t = t.replace('\xa0', ' ')
    return t


def get_fragments(para_element):
    """Get ordered (type, content) fragments from a docx paragraph."""
    fragments = []
    for child in para_element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            rPr = child.find('w:rPr', ns)
            is_bold = rPr is not None and rPr.find('w:b', ns) is not None
            is_italic = rPr is not None and rPr.find('w:i', ns) is not None
            t_el = child.find('w:t', ns)
            if t_el is not None and t_el.text:
                fragments.append(('text', t_el.text, is_bold, is_italic))
        elif tag == 'oMath':
            math_html = omath_to_html(child)
            if math_html:
                fragments.append(('math', math_html, False, True))
        elif tag == 'oMathPara':
            math_html = omath_to_html(child)
            if math_html:
                fragments.append(('displaymath', math_html, False, False))
        elif tag == 'hyperlink':
            for r in child.findall('w:r', ns):
                t_el = r.find('w:t', ns)
                if t_el is not None and t_el.text:
                    fragments.append(('text', t_el.text, False, False))
    return fragments


def rebuild_from_fragments(fragments):
    """Build HTML content from docx fragments."""
    parts = []
    in_bold = False
    in_italic = False

    for frag in fragments:
        ftype, text = frag[0], frag[1]
        is_bold = frag[2] if len(frag) > 2 else False
        is_italic = frag[3] if len(frag) > 3 else False

        if ftype in ('math', 'displaymath'):
            if in_italic: parts.append('</em>'); in_italic = False
            if in_bold: parts.append('</strong>'); in_bold = False
            parts.append(f' <em class="math">{text}</em> ')
        else:
            if is_bold and not in_bold:
                if in_italic: parts.append('</em>'); in_italic = False
                parts.append('<strong>'); in_bold = True
            elif not is_bold and in_bold:
                parts.append('</strong>'); in_bold = False
            if is_italic and not in_italic:
                parts.append('<em>'); in_italic = True
            elif not is_italic and in_italic:
                parts.append('</em>'); in_italic = False
            parts.append(html_mod.escape(text))

    if in_italic: parts.append('</em>')
    if in_bold: parts.append('</strong>')
    return ''.join(parts)


def build_docx_sequence(doc):
    """Build the full sequence of docx paragraphs with their properties."""
    sequence = []
    for p in doc.paragraphs:
        text = p.text.strip()
        has_math = bool(p._element.findall('.//m:oMath', ns))
        has_display = bool(p._element.findall('.//m:oMathPara', ns))
        is_math_only = has_display and not text
        fragments = get_fragments(p._element) if has_math else None

        sequence.append({
            'text': text,
            'normalized': normalize(text) if text else '',
            'has_math': has_math,
            'is_display': has_display,
            'is_math_only': is_math_only,
            'fragments': fragments,
        })
    return sequence


def find_html_line_by_text(lines, search_text, start_from=0):
    """Find the line index in HTML that matches the given text."""
    search_norm = normalize(search_text)
    if len(search_norm) < 8:
        return -1

    for i in range(start_from, len(lines)):
        line = lines[i]
        if '<p ' not in line and '<li' not in line:
            continue
        text_only = normalize(re.sub(r'<[^>]+>', '', line))
        if not text_only:
            continue
        # Check prefix match
        if text_only[:30] == search_norm[:30]:
            return i
        # Check word overlap for shorter texts
        if len(search_norm) > 20:
            s_words = set(search_norm.split())
            t_words = set(text_only.split())
            if s_words and t_words:
                overlap = len(s_words & t_words) / max(len(s_words), len(t_words))
                if overlap > 0.7:
                    return i
    return -1


def process_chapter(filepath, docx_seq, chapter_start, chapter_end):
    """Process a single chapter HTML file using the docx sequence."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    lines = content.split('\n')

    insertions = 0
    rebuilds = 0
    html_cursor = 0

    new_lines = list(lines)  # work on a copy
    offset = 0  # track insertions shifting line numbers

    for di in range(chapter_start, chapter_end):
        entry = docx_seq[di]

        if entry['is_math_only']:
            # This is a display equation - need to INSERT a new paragraph
            # Find the preceding text paragraph in the docx
            prev_text = ''
            for pi in range(di - 1, max(chapter_start - 1, di - 5), -1):
                if docx_seq[pi]['text'] and not docx_seq[pi]['is_math_only']:
                    prev_text = docx_seq[pi]['text']
                    break

            if not prev_text:
                continue

            # Find this preceding paragraph in the HTML
            html_idx = find_html_line_by_text(new_lines, prev_text, html_cursor)
            if html_idx < 0:
                continue

            # Build the display equation HTML
            math_html = omath_to_html(entry['fragments'][0][1]) if entry['fragments'] else ''
            if not math_html:
                # Get math directly from the element
                for p in [docx_seq[di]]:
                    pass
                continue

            eq_line = f'<p class="display-math"><em class="math">{math_html}</em></p>'

            # Insert after the found line
            insert_at = html_idx + 1 + offset
            new_lines.insert(insert_at, eq_line)
            offset += 1
            insertions += 1
            html_cursor = html_idx + 1

        elif entry['has_math'] and entry['fragments']:
            # Paragraph with inline math - check if HTML version is missing math
            html_idx = find_html_line_by_text(new_lines, entry['text'], html_cursor)
            if html_idx < 0:
                continue

            html_line = new_lines[html_idx + offset] if html_idx + offset < len(new_lines) else ''
            text_only = re.sub(r'<[^>]+>', '', html_line)

            # Check if this line still has missing math (double spaces or single-space gaps)
            has_gaps = ('  ' in text_only or
                       re.search(r'\b[A-Za-z] [,.]', text_only) is not None)

            if has_gaps and 'class="math"' not in html_line:
                # Rebuild the paragraph from docx fragments
                tag_match = re.match(r'(<p[^>]*>)', html_line)
                if tag_match:
                    open_tag = tag_match.group(1)
                    rebuilt = rebuild_from_fragments(entry['fragments'])
                    new_lines[html_idx + offset] = f'{open_tag}{rebuilt}</p>'
                    rebuilds += 1

            html_cursor = html_idx + 1

    if insertions > 0 or rebuilds > 0:
        with open(filepath, 'w', encoding='utf-8', newline='\n') as f:
            f.write('\n'.join(new_lines))

    return insertions, rebuilds


def find_chapter_boundaries(docx_seq):
    """Find where each chapter starts/ends in the docx sequence."""
    chapters = {}
    current_chapter = None
    for i, entry in enumerate(docx_seq):
        text = entry['text']
        if text.startswith('Chapter ') and ':' in text[:30]:
            ch_match = re.match(r'Chapter (\d+)', text)
            if ch_match:
                ch_num = int(ch_match.group(1))
                if current_chapter is not None:
                    chapters[current_chapter]['end'] = i
                current_chapter = ch_num
                chapters[ch_num] = {'start': i, 'end': len(docx_seq)}
        elif text.startswith('Appendix '):
            if current_chapter is not None:
                chapters[current_chapter]['end'] = i
                current_chapter = None

    return chapters


# Map chapter numbers to HTML filenames
CHAPTER_FILES = {}
for fname in os.listdir(BOOK_DIR):
    m = re.match(r'chapter-(\d+)-', fname)
    if m:
        CHAPTER_FILES[int(m.group(1))] = fname


def main():
    print("Loading source docx...")
    doc = Document(DOCX_PATH)
    print("Building paragraph sequence...")
    docx_seq = build_docx_sequence(doc)
    print(f"  {len(docx_seq)} total paragraphs")

    # Count display equations
    display_count = sum(1 for e in docx_seq if e['is_math_only'])
    print(f"  {display_count} display equations to insert")

    chapters = find_chapter_boundaries(docx_seq)
    print(f"  {len(chapters)} chapters found\n")

    # For display equations, we need to extract math directly from XML
    # since fragments may not capture oMathPara correctly
    # Re-extract: for each math-only paragraph, get the math HTML
    for i, entry in enumerate(docx_seq):
        if entry['is_math_only']:
            para_el = doc.paragraphs[i]._element
            omath_els = para_el.findall('.//m:oMath', ns)
            if omath_els:
                combined = ' '.join(omath_to_html(om) for om in omath_els)
                entry['display_html'] = combined
            else:
                omathpara_els = para_el.findall('.//m:oMathPara', ns)
                if omathpara_els:
                    combined = ' '.join(omath_to_html(om) for om in omathpara_els)
                    entry['display_html'] = combined

    total_insertions = 0
    total_rebuilds = 0

    for ch_num in sorted(CHAPTER_FILES.keys()):
        if ch_num not in chapters:
            continue

        fname = CHAPTER_FILES[ch_num]
        filepath = os.path.join(BOOK_DIR, fname)
        ch_start = chapters[ch_num]['start']
        ch_end = chapters[ch_num]['end']

        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        lines = content.split('\n')

        insertions = 0
        rebuilds = 0
        new_lines = list(lines)
        offset = 0
        html_cursor = 0

        for di in range(ch_start, ch_end):
            entry = docx_seq[di]

            if entry['is_math_only'] and 'display_html' in entry:
                # Find preceding non-math paragraph
                prev_text = ''
                for pi in range(di - 1, max(ch_start - 1, di - 5), -1):
                    if docx_seq[pi]['text'] and not docx_seq[pi]['is_math_only']:
                        prev_text = docx_seq[pi]['text']
                        break
                if not prev_text:
                    continue

                html_idx = find_html_line_by_text(new_lines, prev_text, html_cursor)
                if html_idx < 0:
                    continue

                # Check if equation already exists after this line
                next_idx = html_idx + 1 + offset
                if next_idx < len(new_lines) and 'display-math' in new_lines[next_idx]:
                    html_cursor = html_idx + 1
                    continue

                eq_html = entry['display_html']
                eq_line = f'<p class="display-math"><em class="math">{eq_html}</em></p>'
                new_lines.insert(html_idx + 1 + offset, eq_line)
                offset += 1
                insertions += 1
                html_cursor = html_idx + 1

            elif entry['has_math'] and entry['text']:
                html_idx = find_html_line_by_text(new_lines, entry['text'], html_cursor)
                if html_idx < 0:
                    continue

                actual_idx = html_idx + offset
                if actual_idx >= len(new_lines):
                    continue
                html_line = new_lines[actual_idx]
                text_only = re.sub(r'<[^>]+>', '', html_line)

                # Check for remaining gaps
                has_gaps = ('  ' in text_only or
                           re.search(r'(?<=[a-z]) [,.]', text_only))

                if has_gaps and 'class="math"' not in html_line and entry['fragments']:
                    tag_match = re.match(r'(<p[^>]*>)', html_line)
                    if tag_match:
                        open_tag = tag_match.group(1)
                        rebuilt = rebuild_from_fragments(entry['fragments'])
                        if 'class="math"' in rebuilt:
                            new_lines[actual_idx] = f'{open_tag}{rebuilt}</p>'
                            rebuilds += 1

                html_cursor = html_idx + 1

        if insertions > 0 or rebuilds > 0:
            with open(filepath, 'w', encoding='utf-8', newline='\n') as f:
                f.write('\n'.join(new_lines))
            print(f"  {fname}: {insertions} equations inserted, {rebuilds} paragraphs rebuilt")
            total_insertions += insertions
            total_rebuilds += rebuilds

    print(f"\nTotal: {total_insertions} display equations inserted, {total_rebuilds} paragraphs rebuilt")

    # Add display-math CSS
    css_path = os.path.join(BOOK_DIR, 'book.css')
    with open(css_path, 'r', encoding='utf-8') as f:
        css = f.read()
    if '.display-math' not in css:
        css += '\n\n/* Display equations */\np.display-math { text-align: center; margin: 1.2em 0; font-size: 1.1em; }\np.display-math em.math { font-style: italic; font-family: "Cambria Math", "STIX Two Math", serif; }\n'
        with open(css_path, 'w', encoding='utf-8') as f:
            f.write(css)
        print("Added .display-math CSS to book.css")


if __name__ == '__main__':
    main()
