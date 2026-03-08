#!/usr/bin/env python3
"""
Build Geometric Ethics v1.17 from v1.16.
Adds Part VI: Domain Applications with five new chapters:
  - Chapter 20: Geometric Economics (BGE, Nash nesting)
  - Chapter 21: Geometric Clinical Ethics
  - Chapter 22: Geometric Jurisprudence
  - Chapter 23: Geometric Finance
  - Chapter 24: Geometric Theology
Renumbers old Part VI -> Part VII, old Ch 20 -> Ch 25, Ch 21 -> Ch 26.
Updates cross-references throughout.
"""

import copy
import re
import sys
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

SRC = Path(r"C:\Users\abptl\Documents\Geometric Ethics - The Mathematical Structure of Moral Reasoning - Bond - v1.16 - Mar 2026.docx")
DST = Path(r"C:\Users\abptl\Documents\Geometric Ethics - The Mathematical Structure of Moral Reasoning - Bond - v1.17 - Mar 2026.docx")

print("Loading v1.16...")
doc = Document(str(SRC))

# ============================================================
# STEP 1: Find insertion point (just before "Part VI: Horizons")
# ============================================================
insert_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and 'Part VI' in p.text and 'Horizon' in p.text:
        insert_idx = i
        break

if insert_idx is None:
    print("ERROR: Could not find 'Part VI: Horizons'")
    sys.exit(1)

print(f"Found Part VI: Horizons at paragraph {insert_idx}")

# ============================================================
# STEP 2: Renumber old Part VI -> Part VII, Ch 20 -> 25, Ch 21 -> 26
# ============================================================
XREF_MAP = {
    'Chapter 20': 'Chapter 25',
    'Chapter 21': 'Chapter 26',
    'Ch. 20': 'Ch. 25',
    'Ch. 21': 'Ch. 26',
    'Ch 20': 'Ch 25',
    'Ch 21': 'Ch 26',
    '\u00a720.': '\u00a725.',
    '\u00a721.': '\u00a726.',
    'Section 20.': 'Section 25.',
    'Section 21.': 'Section 26.',
    'Part VI': 'Part VII',
}

def renumber_text(text):
    """Apply all renumbering to a text string."""
    result = text
    for old, new in XREF_MAP.items():
        result = result.replace(old, new)
    return result

def renumber_runs(paragraph):
    """Renumber cross-references in a paragraph while preserving formatting."""
    for run in paragraph.runs:
        new_text = renumber_text(run.text)
        if new_text != run.text:
            run.text = new_text

print("Renumbering Part VI -> Part VII, Ch 20 -> 25, Ch 21 -> 26...")
renumber_count = 0
for p in doc.paragraphs:
    old_text = p.text
    renumber_runs(p)
    if p.text != old_text:
        renumber_count += 1

print(f"  Renumbered text in {renumber_count} paragraphs")

# Also renumber section numbers in headings of old Ch 20/21
# These appear as "20.1", "20.2", ..., "21.1", "21.2", ... in Heading 2/3 styles
import re
section_renumber_count = 0
for p in doc.paragraphs:
    if p.style.name in ('Heading 2', 'Heading 3', 'Heading2', 'Heading3'):
        for run in p.runs:
            new_text = re.sub(r'\b20\.(\d)', r'25.\1', run.text)
            new_text = re.sub(r'\b21\.(\d)', r'26.\1', new_text)
            if new_text != run.text:
                run.text = new_text
                section_renumber_count += 1

print(f"  Renumbered {section_renumber_count} section headings (20.x -> 25.x, 21.x -> 26.x)")

# ============================================================
# STEP 3: Update version number in header/title
# ============================================================
for p in doc.paragraphs:
    if 'v 1.16' in p.text or 'v1.16' in p.text:
        for run in p.runs:
            run.text = run.text.replace('v 1.16', 'v 1.17').replace('v1.16', 'v1.17')

# Update headers/footers
for section in doc.sections:
    for header in [section.header, section.first_page_header]:
        if header:
            for p in header.paragraphs:
                for run in p.runs:
                    run.text = run.text.replace('v 1.16', 'v 1.17').replace('v1.16', 'v1.17')

print("Updated version to v1.17")

# ============================================================
# STEP 4: Build the new chapter content
# ============================================================

# Import chapter content from separate files
sys.path.insert(0, str(Path(__file__).parent))
from ch21_clinical import CH21_CLINICAL
from ch22_jurisprudence import CH22_JURISPRUDENCE
from ch23_finance import CH23_FINANCE
from ch24_theology import CH24_THEOLOGY

# Part VI intro + Chapter 20 (Economics)
PART_INTRO = [
    ('Heading 1', 'Part VI: Domain Applications'),
    ('Body Text', 'The preceding five parts developed the mathematical framework of Geometric Ethics: the moral manifold, tensor hierarchy, dynamics, symmetry, conservation laws, and implementation architecture. This part demonstrates that the framework is not confined to abstract ethics or AI alignment. The same mathematical structures \u2014 pathfinding on stratified manifolds, gauge invariance, Noether conservation, tensorial contraction \u2014 apply directly to established domains with their own formal traditions: economics, clinical medicine, law, finance, and theology.'),
    ('Body Text', 'Each chapter in this part takes a domain that has struggled with the limitations of scalar models, shows how the geometric framework resolves specific longstanding puzzles, and identifies falsifiable predictions that distinguish the geometric approach from existing domain-specific theories. The chapters can be read independently, but they share a common architecture: domain-specific instantiation of the moral manifold, identification of the relevant dimensions, construction of the domain geodesic, and derivation of results that are inaccessible from scalar models.'),
    ('Body Text', ''),
]

CH20_ECONOMICS = [
    ('Heading 1', 'Chapter 20: Geometric Economics \u2014 The Bond Geodesic Equilibrium'),
    ('Body Text', 'This chapter applies the Geometric Ethics framework to economic decision theory, demonstrating that the classical model of Homo economicus \u2014 a perfectly rational agent maximizing scalar utility \u2014 is not wrong because humans are irrational, but incomplete because it computes on a projected subspace of the actual decision manifold. We construct the Bond Geodesic Equilibrium (BGE), prove it subsumes Nash equilibrium as a special case, and derive prospect-theoretic phenomena as geometric properties of the decision manifold.'),
    ('Body Text', ''),
    ('Heading 2', '20.1 The Failure of Scalar Economics'),
    ('Body Text', 'Classical economics rests on Homo economicus \u2014 the perfectly rational agent who maximizes a scalar utility function over all possible actions. This model is mathematically elegant and empirically false. Behavioral economics, following Kahneman and Tversky, has documented systematic deviations from rational choice but lacks a unified mathematical framework that explains why these deviations occur and what structure they share.'),
    ('Body Text', 'Kahneman\'s dual-process theory \u2014 System 1 (fast, automatic, heuristic) and System 2 (slow, deliberate, calculating) \u2014 provides the psychological architecture. What it lacks is the mathematical content: What is the space on which System 1 and System 2 operate? What is the formal relationship between them? Why does their interaction produce the specific pattern of "biases" that behavioral economics has catalogued?'),
    ('Body Text', 'The answer is that economic decisions are not made on a one-dimensional space. They are made on a multi-dimensional decision manifold that includes monetary cost, moral constraint, social norms, fairness, autonomy, trust, identity, institutional legitimacy, and epistemic status \u2014 the same nine dimensions of the moral manifold M developed in Chapter 5, re-interpreted for economic contexts. Every economic transaction is also a moral event. The classical economist\'s move of stripping away all dimensions except monetary cost is not a simplification \u2014 it is a dimensional collapse that destroys the information needed to predict actual behavior.'),
    ('Body Text', ''),
    ('Heading 2', '20.2 The Economic Decision Complex'),
    ('Body Text', 'Definition 20.1 (Economic Decision Complex). The economic decision complex E is a weighted simplicial complex constructed as follows:'),
    ('Body Text', '\u2022 Vertices (0-simplices): Each vertex v_i represents an economic state \u2014 a configuration of goods, services, obligations, and relationships. The vertex carries an attribute vector a(v_i) \u2208 R^9 scoring the nine dimensions.'),
    ('Body Text', '\u2022 Edges (1-simplices): An edge (v_i, v_j) represents an available action \u2014 a transaction, decision, or behavioral choice. The edge carries a weight w(v_i, v_j) \u2265 0 representing the total cost on the full manifold.'),
    ('Body Text', '\u2022 Higher simplices: A k-simplex [v_0, \u2026, v_k] represents a bundle of mutually available actions.'),
    ('Body Text', 'Definition 20.2 (Multi-Dimensional Edge Weights). The weight of an edge (v_i, v_j) in E is:'),
    ('Body Text', '    w(v_i, v_j) = Da^T Sigma^{-1} Da  +  Sum_k beta_k * 1[moral boundary k crossed]'),
    ('Body Text', 'where Da = a(v_j) - a(v_i) is the attribute-vector difference, Sigma is the 9x9 dimensional covariance matrix estimated from observed economic behavior, and beta_k is the penalty for crossing the k-th moral-economic boundary.'),
    ('Body Text', 'The Mahalanobis distance is critical. Economic dimensions are not independent: fairness (d_3) interacts with consequences (d_1); autonomy (d_4) modulates the weight of social impact (d_6); trust (d_5) gates access to higher-consequence transactions (d_1).'),
    ('Body Text', ''),
    ('Heading 2', '20.3 Heuristic Bounded Rationality'),
    ('Body Text', 'The agent\'s decision on E is A* search (Chapter 11) with evaluation function f(n) = g(n) + h(n), where g(n) is the accumulated economic cost (System 2: slow, deliberate calculation) and h(n) is the moral-heuristic cost (System 1: fast, automatic boundary enforcement).'),
    ('Body Text', 'Theorem 20.1 (Heuristic Truncation). If the moral heuristic h is admissible \u2014 i.e., h(n) \u2264 h*(n) for all nodes n, where h*(n) is the true remaining cost on the full manifold \u2014 then A* search with h finds the optimal path (the Bond geodesic) on E.'),
    ('Body Text', 'Moral heuristics \u2014 "do not steal," "keep promises," "reciprocate kindness" \u2014 are admissible because they never overestimate the true cost of violating a moral boundary. The penalty for stealing (social, legal, psychological) is always at least as large as the heuristic\'s estimate. These are not cognitive biases; they are search heuristics that reduce an intractable optimization to a tractable one, at the cost of provably bounded suboptimality.'),
    ('Body Text', 'This is Simon\'s bounded rationality program, formalized: agents are rational on the manifold they can perceive, using heuristics that provably approximate the optimal solution.'),
    ('Body Text', ''),
    ('Heading 2', '20.4 The Bond Geodesic and the Scalar Irrecoverability Theorem'),
    ('Body Text', 'Definition 20.3 (Bond Geodesic). The Bond geodesic gamma* from economic state s to state t on E is the minimum-cost path:'),
    ('Body Text', '    gamma* = arg min_gamma Sum w(e_i)'),
    ('Body Text', 'where the sum is over edges e_i in path gamma. When boundary penalties are absorbed into the metric, this is a geodesic in the sense of Chapter 10.'),
    ('Body Text', 'Theorem 20.2 (Scalar Irrecoverability). No continuous function phi: R^9 -> R is injective. Therefore, any continuous scalar utility function u = phi(a) applied to the nine-dimensional attribute vector a \u2208 R^9 destroys information that is mathematically irrecoverable.'),
    ('Body Text', 'Proof. By Brouwer\'s invariance of dimension, R^9 and R^1 are not homeomorphic. Restricting to the compact cube [0,1]^9, suppose phi: [0,1]^9 -> R were a continuous injection. Then phi maps a compact space injectively and continuously into a Hausdorff space, so phi is a homeomorphism onto its image \u2014 an embedding of a 9-dimensional space into R^1. This contradicts invariance of dimension. []'),
    ('Body Text', 'This is the Scalar Irrecoverability Theorem of Chapter 15.6, applied to the economic domain. "Irrational" behavior is rational on the full manifold, irrational only on the projection. Homo economicus computes on the wrong manifold.'),
    ('Body Text', ''),
    ('Heading 2', '20.5 The Bond Geodesic Equilibrium'),
    ('Body Text', 'When multiple agents pathfind simultaneously on the economic decision complex, we enter game theory. The Bond Geodesic Equilibrium (BGE) is the natural equilibrium concept on the manifold.'),
    ('Body Text', 'Definition 20.4 (Bond Geodesic Equilibrium). A strategy profile (gamma_1*, \u2026, gamma_n*) is a Bond Geodesic Equilibrium if for each agent i, the path gamma_i* is a Bond geodesic on agent i\'s perceived decision complex, given the paths chosen by all other agents.'),
    ('Body Text', 'To establish the relationship between BGE and Nash equilibrium, we construct the augmented game.'),
    ('Body Text', 'Definition 20.5 (Augmented Game). Given a multi-agent decision problem on E, the augmented game is Gamma+ = (N, {S_i}, {u_i}) where:'),
    ('Body Text', '\u2022 N is the set of agents'),
    ('Body Text', '\u2022 S_i is agent i\'s set of available paths on E'),
    ('Body Text', '\u2022 u_i = -BF_i is agent i\'s payoff: the negative of the total behavioral friction (edge weight sum) along their chosen path, given others\' paths'),
    ('Body Text', 'This construction is canonical: it says that agents in the manifold game are maximizing the negative of their total cost on the full manifold \u2014 i.e., they are minimizing behavioral friction, which is exactly what A* search does.'),
    ('Body Text', ''),
    ('Heading 3', 'The BGE\u2013Nash Relationship'),
    ('Body Text', 'Theorem 20.3 (BGE\u2013Nash Relationship).'),
    ('Body Text', '(1) Equivalence: A strategy profile is a BGE of the manifold game if and only if it is a Nash equilibrium of the augmented game Gamma+.'),
    ('Body Text', '(2) Scalar nesting: If all non-monetary dimensions are zeroed out (i.e., the attribute vector collapses to a(v) = (d_1(v), 0, \u2026, 0)) and each agent selects a single action (not a path), then BGE reduces to Nash equilibrium of the classical game with payoffs u_i = d_1.'),
    ('Body Text', '(3) Refinement: In games where agents can operate on the full manifold, BGE selects among Nash equilibria by the manifold structure \u2014 the equilibrium that minimizes total behavioral friction across all dimensions.'),
    ('Body Text', 'Proof. (1) By construction, gamma_i* is a Bond geodesic if and only if it minimizes BF_i given others\' paths, which is exactly the Nash condition on Gamma+. (2) When a = (d_1, 0, \u2026, 0) and paths are single edges, BF_i = d_1 cost, so -BF_i = -d_1 cost = classical payoff (up to sign). (3) Among multiple Nash equilibria of Gamma+, the BGE selects the profile with minimum total BF, since that is what A* search converges to. []'),
    ('Body Text', 'This is the central result: Nash equilibrium is the d_1-only projection of BGE, exactly as Newtonian mechanics is the low-velocity limit of general relativity. The projection is the contraction of Chapter 15 applied to game theory.'),
    ('Body Text', ''),
    ('Heading 3', 'Existence and Uniqueness'),
    ('Body Text', 'Theorem 20.4 (Existence of Mixed BGE). Every finite augmented game Gamma+ has at least one mixed BGE.'),
    ('Body Text', 'Proof. By Nash\'s theorem (1950), every finite game has a mixed-strategy Nash equilibrium. Gamma+ is a finite game (finite agents, finite paths). By Theorem 20.3(1), a mixed Nash equilibrium of Gamma+ is a mixed BGE. []'),
    ('Body Text', 'Definition 20.6 (Mixed BGE). A mixed BGE is a profile of probability distributions (sigma_1, \u2026, sigma_n) over paths, where each sigma_i minimizes agent i\'s expected behavioral friction given others\' mixed strategies.'),
    ('Body Text', 'Remark (Pure BGE Existence). Pure BGE need not exist in general (as with pure Nash). Sufficient conditions include: (a) the augmented game is a potential game; (b) the augmented game has supermodular structure; (c) agents\' decision complexes are independent (no cross-agent edge weight dependencies).'),
    ('Body Text', ''),
    ('Heading 3', 'The Contraction Lemma'),
    ('Body Text', 'Lemma 20.1 (Contraction Condition for BGE Uniqueness). For each agent i, let alpha_i > 0 measure the self-separation of agent i\'s best-response mapping (how much i\'s own behavioral friction changes with i\'s own path choice) and let kappa_i >= 0 measure the cross-sensitivity (how much i\'s behavioral friction changes with others\' path choices). If'),
    ('Body Text', '    alpha_i > kappa_i    for all i in N'),
    ('Body Text', 'then the best-response mapping on Gamma+ is a contraction with Lipschitz constant L = max_i(kappa_i / alpha_i) < 1, and the BGE is unique.'),
    ('Body Text', 'Proof. The best-response mapping T: S -> S maps each profile to the profile of individual best responses. For agent i, ||T_i(s) - T_i(s\')||\u2264 (kappa_i/alpha_i)||s_{-i} - s\'_{-i}|| by the ratio of cross-sensitivity to self-separation. Taking the maximum over i, ||T(s) - T(s\')||\u2264 L||s - s\'|| with L = max_i(kappa_i/alpha_i) < 1. By Banach\'s fixed-point theorem, T has a unique fixed point. []'),
    ('Body Text', 'Interpretation: BGE is unique when each agent\'s decision is dominated by their own manifold structure (high alpha_i) rather than by what others do (low kappa_i). This is the "weak coupling" condition. When coupling is strong (kappa_i \u2248 alpha_i), multiple BGE can coexist \u2014 as in coordination games, where the equilibrium depends on which manifold region the agents coordinate on.'),
    ('Body Text', ''),
    ('Heading 2', '20.6 Welfare Properties'),
    ('Body Text', 'Theorem 20.5 (BGE Welfare). A BGE is Nash-optimal on the augmented game Gamma+: no agent can unilaterally reduce their behavioral friction by deviating.'),
    ('Body Text', 'However, a BGE is not in general Pareto optimal. Consider the Prisoner\'s Dilemma on the manifold: if the fairness dimension d_3 is inactive, both agents defect (as in classical Nash); the mutual-cooperation profile has lower total BF but is not individually stable. The Pareto failure of Nash equilibrium persists in BGE for the same structural reason.'),
    ('Body Text', 'Remark. In potential games \u2014 where all agents\' behavioral friction derives from a common potential function \u2014 the BGE does coincide with a Pareto optimum of the augmented game. This is a strong sufficient condition that holds in many economically important settings (congestion games, market entry games).'),
    ('Body Text', ''),
    ('Heading 2', '20.7 Behavioral Game Theory as Manifold Geometry'),
    ('Body Text', 'The BGE framework resolves several longstanding puzzles in behavioral game theory by identifying the active manifold dimensions:'),
    ('Body Text', 'Ultimatum game rejections. In the standard ultimatum game, a proposer offers a split of $10; the responder can accept or reject (destroying both payoffs). Homo economicus predicts the responder accepts any positive offer. Experimentally, offers below ~30% are rejected roughly half the time. The BGE explanation: the responder\'s decision complex includes d_3 (fairness). Rejecting an unfair offer has high monetary cost (d_1) but avoids a large fairness boundary penalty. The Bond geodesic includes rejection when the fairness penalty exceeds the monetary cost.'),
    ('Body Text', 'Trust game cooperation. In trust games, agents cooperate far more than Nash predicts. The BGE explanation: dimension d_5 (trust) is active. Cooperation is the Bond geodesic when the trust dimension\'s edge weights make defection more costly on the full manifold than on the d_1 projection.'),
    ('Body Text', 'Public goods provision. Agents contribute to public goods even when free-riding is the Nash strategy. The BGE explanation: dimension d_6 (social impact) creates boundary penalties for free-riding that exceed the monetary savings.'),
    ('Body Text', 'Each of these "anomalies" is predicted by BGE, not catalogued post hoc. The framework says: tell me the active dimensions and their weights, and I will tell you the equilibrium. This is a stronger claim than behavioral economics makes, and it is falsifiable.'),
    ('Body Text', ''),
    ('Heading 2', '20.8 Prospect Theory as Manifold Geometry'),
    ('Body Text', 'The major phenomena of Kahneman and Tversky\'s prospect theory are derived as geometric properties of the decision manifold:'),
    ('Body Text', 'Loss aversion. The edge weights in E are asymmetric: the weight for moving from state a to a worse state b (loss) exceeds the weight for the reverse move (gain), even when |Da| is the same. Formally, w(a -> b) > w(b -> a) when a is the reference point. This is not a "bias" \u2014 it is an asymmetry of the manifold metric, analogous to the asymmetry of a Finsler metric.'),
    ('Body Text', 'Framing effects. A framing effect \u2014 where logically equivalent descriptions of the same decision produce different choices \u2014 is a gauge symmetry violation on the decision manifold. The Bond Invariance Principle (BIP, Chapter 12) requires that evaluations be invariant under meaning-preserving transformations. Framing effects occur when the agent\'s heuristic function h(n) is not gauge-invariant. This is a miscalibration of the heuristic, not a fundamental feature of the manifold.'),
    ('Body Text', 'Reference dependence. The edge weights depend on the starting vertex \u2014 the agent\'s current state serves as the reference point. This is built into the Mahalanobis distance: Da = a(v_j) - a(v_i) is measured from the current state, not from an absolute origin.'),
    ('Body Text', 'Hyperbolic discounting. The temporal dimension of the manifold has non-constant curvature: near-future and far-future states are not metrically equivalent. The "present bias" is a consequence of higher edge-weight density near the current temporal position.'),
    ('Body Text', 'The endowment effect. Owning a good changes the agent\'s position on the manifold (the attribute vector shifts). The Mahalanobis distance from "own the good" to "sell the good" differs from the distance from "don\'t own" to "buy the good" because the reference point has changed. This is not irrationality \u2014 it is path-dependence on a curved manifold.'),
    ('Body Text', ''),
    ('Heading 2', '20.9 Attribute Conservation in Bilateral Exchange'),
    ('Body Text', 'Theorem 20.6 (Attribute Conservation). In a closed bilateral exchange between agents A and B, for each transferable dimension k \u2208 {d_1 (monetary value), d_2 (rights/entitlements), d_4 (autonomy)}:'),
    ('Body Text', '    Da_k(A) + Da_k(B) = 0'),
    ('Body Text', 'That is, what A gains on dimension k, B loses, and vice versa.'),
    ('Body Text', 'This is the economic specialization of the Conservation of Harm (Chapter 12, Theorem 12.1): in a closed system with re-description invariance, conserved quantities cannot be created or destroyed by relabeling. For transferable dimensions \u2014 money, property rights, autonomy \u2014 bilateral exchange is zero-sum on each dimension separately.'),
    ('Body Text', 'Remark (Non-Conservation of Evaluative Dimensions). Evaluative dimensions \u2014 d_3 (fairness), d_7 (identity/virtue) \u2014 are not conserved. Both parties to an exchange can simultaneously perceive increased fairness, or both can perceive unfairness. A fair exchange creates mutual value on d_3 that did not exist before the exchange. This distinction between transferable and evaluative dimensions formalizes the intuition that trade creates value beyond the material transfer.'),
    ('Body Text', ''),
    ('Heading 2', '20.10 Falsifiable Predictions'),
    ('Body Text', 'The framework generates six predictions that distinguish it from both classical and behavioral economics:'),
    ('Body Text', 'Prediction 1 (Dimensional Activation): Ultimatum game rejection rates should vary with the salience of specific manifold dimensions. Making fairness (d_3) more salient (e.g., by framing the game explicitly as a fairness test) should increase rejections; making it less salient (anonymous, one-shot, large stakes) should decrease them. What would falsify: if rejection rates are invariant to dimensional salience manipulations.'),
    ('Body Text', 'Prediction 2 (Bond Index Correlates): Agents with higher Bond Index scores should (a) reject more unfair ultimatum offers, (b) contribute more in public goods games, (c) show higher WTP for fair-trade goods, and (d) exhibit smaller framing effects. What would falsify: if Bond Index is uncorrelated with these measures.'),
    ('Body Text', 'Prediction 3 (Cross-Cultural Metric Variation): The covariance matrix Sigma should vary systematically across cultures in ways predicted by known cultural dimensions. What would falsify: if Sigma is invariant across cultures, or varies unsystematically.'),
    ('Body Text', 'Prediction 4 (Boundary Penalty Measurement): Sacred-value boundaries (beta_k = infinity) should produce qualitatively different response patterns from high-but-finite boundaries. What would falsify: if all boundary effects are graded.'),
    ('Body Text', 'Prediction 5 (Heuristic Admissibility): Moral heuristics should satisfy admissibility: h(n) \u2264 h*(n). Agents using cultural heuristics should find paths that are suboptimal but within epsilon of optimal. What would falsify: if moral heuristics systematically overestimate costs.'),
    ('Body Text', 'Prediction 6 (Manifold Dimensionality): Factor analysis of economic behavioral data should recover approximately nine independent dimensions. What would falsify: if the factor structure is consistently lower-rank (fewer than seven) or higher-rank (more than eleven).'),
    ('Body Text', ''),
    ('Heading 2', '20.11 Connection to the Framework'),
    ('Body Text', 'The Bond Geodesic Equilibrium completes the multi-agent story that began in Chapter 14 (Collective Moral Agency):'),
    ('Body Text', '\u2022 Chapter 11 solved single-agent pathfinding: f(n) = g(n) + h(n) on the moral manifold.'),
    ('Body Text', '\u2022 Chapter 14 built the collective agency tensor: how multi-agent moral structure is represented.'),
    ('Body Text', '\u2022 This chapter answers: what is the equilibrium when multiple agents pathfind simultaneously?'),
    ('Body Text', '\u2022 Chapter 15 asks: how does this equilibrium get contracted to a scalar decision?'),
    ('Body Text', 'The BGE is the game-theoretic completion of the geometric ethics program. Nash equilibrium \u2014 the foundation of modern economics \u2014 is the d_1-only contraction of BGE, exactly as scalar utility is the rank-0 contraction of the moral tensor (Chapter 6.7). The same mathematical operation (contraction, Chapter 15) that explains why scalar ethics loses information also explains why scalar economics loses information.'),
    ('Body Text', ''),
    ('Heading 2', '20.12 Summary'),
    ('Body Text', 'This chapter has shown that the geometric ethics framework, when applied to economic decision theory, yields:'),
    ('Body Text', '1. A formal construction of the economic decision complex E as a domain-specific instantiation of the moral manifold M.'),
    ('Body Text', '2. Heuristic bounded rationality: Simon\'s program formalized via A* search.'),
    ('Body Text', '3. The Bond Geodesic Equilibrium: a generalization of Nash equilibrium with existence via Nash\'s theorem on the augmented game, uniqueness via the contraction lemma, and a formal nesting of Nash as the scalar special case.'),
    ('Body Text', '4. Geometric behavioral economics: prospect-theoretic phenomena derived as properties of the manifold metric.'),
    ('Body Text', '5. Attribute conservation: transferable dimensions are conserved in bilateral exchange; evaluative dimensions permit mutual value creation.'),
    ('Body Text', '6. Six falsifiable predictions distinguishing the framework from classical and behavioral economics.'),
    ('Body Text', 'The next four chapters extend this pattern to clinical medicine, law, financial markets, and theology.'),
    ('Body Text', ''),
]

# Assemble all domain chapters
ALL_CHAPTERS = PART_INTRO + CH20_ECONOMICS + CH21_CLINICAL + CH22_JURISPRUDENCE + CH23_FINANCE + CH24_THEOLOGY

# ============================================================
# STEP 5: Insert all content before Part VII (was Part VI)
# ============================================================
print(f"Inserting new Part VI: Domain Applications ({len(ALL_CHAPTERS)} paragraphs) before paragraph {insert_idx}...")

target_element = doc.paragraphs[insert_idx]._element

from lxml import etree
from docx.oxml.ns import qn

for style, text in ALL_CHAPTERS:
    new_para = etree.SubElement(target_element.getparent(), qn('w:p'))
    new_para.getparent().remove(new_para)

    new_para = copy.deepcopy(doc.paragraphs[0]._element)
    for child in list(new_para):
        new_para.remove(child)

    pPr = etree.SubElement(new_para, qn('w:pPr'))
    pStyle = etree.SubElement(pPr, qn('w:pStyle'))
    style_val = style.replace(' ', '')
    pStyle.set(qn('w:val'), style_val)

    run_elem = etree.SubElement(new_para, qn('w:r'))
    t_elem = etree.SubElement(run_elem, qn('w:t'))
    t_elem.text = text
    t_elem.set(qn('xml:space'), 'preserve')

    target_element.addprevious(new_para)

print(f"Inserted {len(ALL_CHAPTERS)} paragraphs")

# ============================================================
# STEP 6: Save
# ============================================================
print(f"Saving to {DST}...")
doc.save(str(DST))
print("Done! v1.17 saved.")
print()
print("=== MANUAL STEPS REQUIRED ===")
print("1. Open v1.17 in Word and regenerate the Table of Contents")
print("2. Update the title page version number if not caught by script")
print("3. Review cross-references: all 'Chapter 20/21' refs now point to 'Chapter 25/26'")
print("4. Add new chapters to the Index")
print("5. Update Core Objects at a Glance to include BGE, Clinical Geodesic, Judicial Complex, Financial Decision Complex")
print("6. Update Key Results at a Glance to include BGE-Nash nesting, QALY Irrecoverability, Topological Constitutionality")
print("7. Update Preface to mention Part VI: Domain Applications")
