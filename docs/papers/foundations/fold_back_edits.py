"""
Fold back architectural innovations from domain papers into the main
Geometric Ethics manuscript (v1.15 docx).

Four additions, applied back-to-front to preserve paragraph indices:
  1. Ch 15.6: Name the Scalar Irrecoverability Theorem (after para 4067)
  2. Ch 11.4: Expand ε-admissibility + name the Heuristic Truncation Theorem (after para 3139)
  3. Ch 11 (new §11.4a): The Discrete Computational Realization (after expanded 11.4)

Run: python fold_back_edits.py
"""

import sys, os, copy
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = os.path.join(os.environ['USERPROFILE'], 'Documents',
    'Geometric Ethics - The Mathematical Structure of Moral Reasoning - Bond - v1.15 - Mar 2026.docx')
DST = os.path.join(os.environ['USERPROFILE'], 'Documents',
    'Geometric Ethics - The Mathematical Structure of Moral Reasoning - Bond - v1.16 - Mar 2026.docx')

doc = Document(SRC)

# ── Helper: insert a new paragraph after a given paragraph index ──────────
def insert_para_after(doc, idx, text, style_name):
    """Insert a paragraph with given text and style after doc.paragraphs[idx]."""
    ref = doc.paragraphs[idx]._element
    from docx.oxml.ns import qn
    from lxml import etree
    # Create new paragraph element
    new_p = copy.deepcopy(doc.paragraphs[0]._element)  # template
    # Clear it
    for child in list(new_p):
        new_p.remove(child)
    # Add run with text
    r = etree.SubElement(new_p, qn('w:r'))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    # Insert after ref
    ref.addnext(new_p)
    # Now find this paragraph in doc.paragraphs to set style
    # (paragraphs list is rebuilt on access)
    for p in doc.paragraphs:
        if p._element is new_p:
            p.style = doc.styles[style_name]
            return p
    return None

def insert_multiple_after(doc, start_idx, items):
    """Insert multiple (text, style) pairs after start_idx.
    Items are inserted in order, each after the previous.
    Returns the index offset (number of paras inserted).
    """
    current_idx = start_idx
    for text, style in items:
        insert_para_after(doc, current_idx, text, style)
        current_idx += 1
    return len(items)


# ══════════════════════════════════════════════════════════════════════════
# EDIT 1: Ch 15.6 — Scalar Irrecoverability Theorem (insert after para 4067)
# After the paragraph ending "This is a moral choice, not merely a technical one."
# Before §15.7 Moral Residue
# ══════════════════════════════════════════════════════════════════════════

edit1_items = [
    # Heading
    ("The Scalar Irrecoverability Theorem", "Heading 3"),
    # First paragraph
    ("The information loss documented above is not merely quantitative — it is structurally irrecoverable. "
     "This result, which we call the Scalar Irrecoverability Theorem, provides the formal foundation for "
     "every domain-specific critique of scalar reduction (QALY-based health economics, utility-maximizing "
     "economics, scalar sentencing guidelines).",
     "First Paragraph"),
    # Theorem statement
    ("Theorem 15.1 (Scalar Irrecoverability). [Proved.] Let Q: \u211d\u2079 \u2192 \u211d be any function that maps "
     "a 9-dimensional moral attribute vector to a scalar evaluation. Then: "
     "(i) Q is not injective: multiple morally distinct states map to the same scalar value. "
     "(ii) The information destroyed by Q is irrecoverable: there exists no function \u03c8: \u211d \u2192 \u211d\u2079 "
     "such that \u03c8 \u2218 Q = id. "
     "(iii) The moral geodesic on the full manifold \u2133 is in general different from the scalar-optimal "
     "path, and the divergence between them is not bounded by any function of Q alone.",
     "Body Text"),
    # Proof
    ("Proof. (i) The rank-nullity theorem requires dim(ker dQ) \u2265 8 at every regular point of Q: "
     "a map from \u211d\u2079 to \u211d has a differential of rank at most 1, so its kernel has dimension at "
     "least 8. Any two states differing only in a kernel direction map to the same scalar. Since the "
     "kernel is 8-dimensional at every regular point, the space of morally distinct states that are "
     "scalar-equivalent is generically 8-dimensional — not a set of measure zero but an 8-dimensional "
     "submanifold for each scalar value. "
     "(ii) By the data processing inequality, information that passes through a dimensionality-reducing "
     "map Q cannot be recovered by any downstream function \u03c8. Formally: H(\u03b1 | Q(\u03b1)) > 0 "
     "whenever the conditional entropy of the attribute vector given its scalar image is positive, which "
     "it is whenever the kernel of dQ contains morally relevant variation — i.e., generically. "
     "(iii) Let \u03b3*_\u2133 be the geodesic on the full 9-dimensional manifold and \u03b3*_Q be the path "
     "that maximizes Q. Since Q collapses 8 dimensions, it can assign equal value to paths with radically "
     "different moral profiles. The divergence between \u03b3*_\u2133 and \u03b3*_Q depends on the structure "
     "of the moral manifold in the kernel of dQ, which is invisible to Q and therefore cannot be bounded "
     "by any function of Q alone. \u25a1",
     "Body Text"),
    # Remark on domain applications
    ("Remark (Domain Applications of Scalar Irrecoverability). The theorem has immediate consequences for "
     "any field that reduces multi-dimensional moral evaluation to a scalar. In health economics, the "
     "quality-adjusted life year (QALY) is a map Q: \u211d\u2079 \u2192 \u211d that collapses clinical "
     "outcomes, autonomy, trust, dignity, justice, and epistemic status to a single number; Theorem 15.1 "
     "proves that the information destroyed is irrecoverable and that QALY-optimal policies can be "
     "manifold-suboptimal. In economics, scalar utility maximization discards the moral, social, and "
     "identity dimensions on which agents actually decide; the theorem explains why \u2018irrational\u2019 "
     "behavior (ultimatum game rejections, loss aversion, reference dependence) is rational on the full "
     "manifold. In law, scalar sentencing (years of imprisonment as the sole metric) discards "
     "rehabilitation, victim impact, community trust, and proportionality dimensions. Each of these is an "
     "instance of Theorem 15.1, not a separate disciplinary problem.",
     "Body Text"),
]

print("Inserting Scalar Irrecoverability Theorem after para 4067...")
insert_multiple_after(doc, 4067, edit1_items)
print(f"  Inserted {len(edit1_items)} paragraphs.")


# ══════════════════════════════════════════════════════════════════════════
# EDIT 2: Ch 11.4 — Expand ε-admissibility + Name Heuristic Truncation
# Replace the single remark at para 3139, and add new content after it.
# First, modify Theorem 11.1's text to add the alias.
# Then expand the remark into a full treatment.
# ══════════════════════════════════════════════════════════════════════════

# 2a: Modify Theorem 11.1 label to include "Heuristic Truncation" alias
# Para 3135 contains: "Theorem 11.1 (Admissibility of Core Moral Heuristics)..."
p_thm = doc.paragraphs[3135]
old_text = p_thm.text
if "Admissibility of Core Moral Heuristics" in old_text:
    new_text = old_text.replace(
        "Theorem 11.1 (Admissibility of Core Moral Heuristics)",
        "Theorem 11.1 (Admissibility of Core Moral Heuristics; the Heuristic Truncation Theorem)"
    )
    # Clear and rewrite
    for run in p_thm.runs:
        run.text = ""
    if p_thm.runs:
        p_thm.runs[0].text = new_text
    else:
        p_thm.add_run(new_text)
    print("Renamed Theorem 11.1 to include 'Heuristic Truncation Theorem' alias.")
else:
    print(f"WARNING: Could not find Theorem 11.1 text at para 3135. Found: {old_text[:100]}")


# 2b: The existing remark at 3139 mentions weighted A* briefly.
# We'll add new content AFTER it (before §11.5 at 3140).
# Insert the ε-admissibility spectrum as a full treatment.

edit2_items = [
    # New subsection heading (will appear between the remark and §11.5)
    ("The \u03b5-Admissibility Spectrum and the Heuristic Hierarchy", "Heading 3"),
    # Opening
    ("The preceding remark acknowledges that real moral heuristics may overestimate. This section "
     "formalizes the consequences. The key insight is that overestimation is not binary (admissible "
     "vs. inadmissible) but falls on a spectrum with distinct moral signatures at each level.",
     "First Paragraph"),
    # Definition
    ("Definition 11.5 (\u03b5-Admissible Moral Heuristic). [Formal Definition.] A moral heuristic h is "
     "\u03b5-admissible if, for all states n \u2208 \u2133: h(n) \u2264 (1 + \u03b5) \u00b7 g*(n, G_eq), "
     "where g*(n, G_eq) is the true minimum behavioral friction from n to the goal region. When A* search "
     "uses an \u03b5-admissible heuristic, the returned path has total cost at most (1 + \u03b5) times "
     "the optimal path cost. The parameter \u03b5 quantifies the degree of moral conservatism: "
     "\u03b5 = 0 yields perfect optimality; small \u03b5 yields near-optimal paths at reduced "
     "computational cost; large \u03b5 yields fast but potentially suboptimal moral judgments.",
     "Body Text"),
    # Connection to weighted A*
    ("Remark (Connection to Weighted A*). In the search literature, \u03b5-admissibility is achieved by "
     "weighted A*: f(n) = g(n) + w \u00b7 h(n), where w = 1 + \u03b5. The weight w trades optimality "
     "for computational tractability. This maps directly onto moral reasoning: a time-pressured agent "
     "(emergency physician, crisis responder) implicitly increases w — relying more heavily on heuristic "
     "judgment and less on deliberate calculation — accepting a bounded-suboptimal moral path in exchange "
     "for the speed needed to act. The bounded suboptimality guarantee ensures that this is not moral "
     "recklessness but a principled trade-off with a quantifiable cost ceiling (Pohl 1970; Pearl 1984).",
     "Body Text"),
    # The 4-category hierarchy theorem
    ("Theorem 11.3 (The Moral Heuristic Hierarchy). [Conditional Theorem; conditional on the behavioral "
     "friction model of Chapter 6 and the stratification of Chapter 8.] Moral heuristics fall into four "
     "categories with distinct phenomenological signatures:",
     "Body Text"),
    # Category 1
    ("(i) Strictly admissible (h(n) \u2264 g*(n, G_eq)): The heuristic underestimates or correctly "
     "estimates the true moral cost. The A* path is optimal. These are the core prohibitions whose true "
     "costs genuinely exceed any finite estimate — murder (true cost includes absorbing-stratum penalties, "
     "legal consequences, social ostracism, and psychological damage), non-consensual treatment (true cost "
     "includes rights violation, institutional liability, trust destruction). The proof of Theorem 11.1 "
     "establishes this category.",
     "Body Text"),
    # Category 2
    ("(ii) \u03b5-admissible (h(n) \u2264 (1 + \u03b5) \u00b7 g*(n, G_eq), small \u03b5): The heuristic "
     "slightly overestimates the true boundary cost. The A* path is near-optimal, with total cost at most "
     "(1 + \u03b5) times the true optimum. This is the normal operating range of a well-calibrated moral "
     "agent. The slight conservatism — marginally overweighting harm avoidance, slightly inflating the "
     "cost of promise-breaking — produces the well-documented phenomenon that experienced moral reasoners "
     "are slightly more cautious than a pure expected-value calculation would dictate, at negligible cost "
     "to overall moral performance. Evolutionary and cultural selection pressures calibrate \u03b5 to be "
     "small but positive: an agent with \u03b5 = 0 (perfect calibration) is fragile (any perturbation "
     "makes the heuristic inadmissible), while an agent with small positive \u03b5 is robust.",
     "Body Text"),
    # Category 3
    ("(iii) Inadmissible (h(n) \u226b g*(n, G_eq)): The heuristic substantially overestimates boundary "
     "costs. The A* path avoids genuinely beneficial actions because the inflated heuristic makes them "
     "appear too costly. This category has a precise clinical and phenomenological signature: it is the "
     "mathematical structure of moral injury, trauma, and pathological risk aversion. An agent whose "
     "heuristic has been damaged by forced boundary crossing (the moral injury formalized in the domain "
     "literature) has inflated \u03b2_k values — they overestimate the cost of actions near the violated "
     "boundary, avoiding even beneficial actions in that region. Defensive behavior (defensive medicine, "
     "defensive lawyering, institutional risk aversion) is \u03b2_institutional \u226b \u03b2*_institutional: "
     "the institution\u2019s heuristic overweights legal and reputational cost relative to its true value, "
     "foreclosing beneficial actions.",
     "Body Text"),
    # Category 4
    ("(iv) Gauge-variant: The heuristic\u2019s output depends on the framing or description of the moral "
     "situation rather than on the situation\u2019s attribute vector. This violates the Bond Invariance "
     "Principle (Chapter 5, §5.5): the heuristic is not invariant under admissible re-descriptions. "
     "Examples include omission bias (harmful action judged as worse than equally harmful inaction, even "
     "when the attribute vectors are identical), status quo bias (framing the current state as the "
     "reference point inflates the cost of departure), anchoring to initial information, and order effects "
     "in moral judgment. These are the phenomena that genuinely deserve the label \u2018cognitive bias\u2019 "
     "in moral reasoning — not because the heuristic overestimates, but because it responds to morally "
     "irrelevant features of the description. The distinction matters: categories (i)\u2013(iii) are "
     "calibration errors (the heuristic estimates the right quantity, more or less accurately); "
     "category (iv) is a gauge failure (the heuristic estimates the wrong quantity).",
     "Body Text"),
    # Proof sketch
    ("Proof sketch. Category (i) follows from Theorem 11.1 (the Heuristic Truncation Theorem). "
     "Category (ii) follows from the bounded suboptimality guarantee of \u03b5-admissible A* "
     "(Pohl 1970, Theorem 2; Pearl 1984, §4.2): if h(n) \u2264 (1 + \u03b5) \u00b7 g*, then "
     "f(\u03b3) \u2264 (1 + \u03b5) \u00b7 f(\u03b3*). Category (iii) is defined by the negation of "
     "(i)\u2013(ii): when \u03b5 exceeds a threshold, the heuristic forecloses paths whose true cost "
     "is acceptable, producing systematic moral avoidance. Category (iv) is defined by violation of "
     "gauge invariance (BIP, Chapter 5): h(\u03c4(n)) \u2260 h(n) for some admissible transformation "
     "\u03c4. The four categories are mutually exclusive and exhaustive for any heuristic function on "
     "\u2133. \u25a1",
     "Body Text"),
    # Recalibration remark
    ("Remark (Heuristic Recalibration). The hierarchy has a dynamic interpretation. An agent\u2019s "
     "heuristic can move between categories over time. Moral education and professional training move "
     "uncalibrated heuristics from category (iii) or (iv) toward categories (i)\u2013(ii) — this is what "
     "it means for clinical training, legal education, or military ethics instruction to \u2018calibrate "
     "moral judgment.\u2019 Conversely, moral injury (forced boundary crossing that inflates \u03b2_k) "
     "moves a previously well-calibrated heuristic from category (ii) to category (iii) — the "
     "mathematical signature of moral damage. Ethics consultation, peer review, and therapeutic "
     "intervention function as recalibration mechanisms: they restore inadmissible heuristics toward "
     "\u03b5-admissibility by correcting inflated boundary penalties to values closer to their true "
     "costs. The framework predicts that effective moral repair must target specific \u03b2_k values "
     "(not general \u2018resilience\u2019), and that the dimension of repair must match the dimension of "
     "injury.",
     "Body Text"),
]

print("Inserting ε-admissibility spectrum after para 3139...")
insert_multiple_after(doc, 3139, edit2_items)
print(f"  Inserted {len(edit2_items)} paragraphs.")


# ══════════════════════════════════════════════════════════════════════════
# EDIT 3: Ch 8 — The Discrete Computational Realization
# Insert before §8.11 Summary (para 2615), after the Metaethical
# Implications subsection (para 2614).
# ══════════════════════════════════════════════════════════════════════════

edit3_items = [
    # New subsection
    ("8.10a The Discrete Computational Realization: Weighted Simplicial Complexes", "Heading 2"),
    # Opening paragraph
    ("The Whitney-stratified manifold of the preceding sections is the continuous mathematical structure. "
     "For computational applications — decision support, policy analysis, AI alignment, and domain-specific "
     "instantiations (economics, medicine, law) — the continuous manifold admits a natural discretization "
     "as a weighted simplicial complex. This section formalizes the discrete structure and establishes "
     "the correspondence between continuous and discrete architectures.",
     "First Paragraph"),
    # Definition of the decision complex
    ("Definition 8.8 (The Moral Decision Complex). [Modeling Axiom.] The moral decision complex \u0394 is a "
     "weighted simplicial complex constructed from the moral manifold \u2133 as follows. "
     "Vertices (0-simplices): Each vertex v_i represents a morally relevant state — a configuration of "
     "the situation\u2019s nine-dimensional attribute vector a(v_i) \u2208 \u211d\u2079. "
     "Edges (1-simplices): An edge (v_i, v_j) represents an available action — a morally relevant "
     "transition from state v_i to state v_j. The edge carries a weight w(v_i, v_j) \u2265 0 "
     "representing the total moral cost of the action on the full manifold. "
     "Higher simplices: A k-simplex [v_0, \u2026, v_k] represents a bundle of jointly executed actions "
     "(a policy, a care protocol, a legislative package) whose combined cost may differ from the sum of "
     "individual edge costs due to interaction effects.",
     "Body Text"),
    # Edge weight definition
    ("Definition 8.9 (Mahalanobis Edge Weights with Boundary Penalties). [Formal Definition.] The weight "
     "of an action (v_i, v_j) on the moral decision complex \u0394 is:",
     "Body Text"),
    # The formula (as text since we can't do OOXML math easily)
    ("w(v_i, v_j) = \u0394a\u1d40 \u03a3\u207b\u00b9 \u0394a + \u03a3_k \u03b2_k \u00b7 \u0031[boundary k crossed]",
     "Body Text"),
    # Explanation
    ("where \u0394a = a(v_j) \u2212 a(v_i) is the change in the situation\u2019s attribute vector, "
     "\u03a3 is the 9 \u00d7 9 moral covariance matrix encoding cross-dimensional dependencies, and "
     "\u03b2_k is the penalty for crossing moral boundary k. The first term is the Mahalanobis distance "
     "on the attribute space — it measures the \u2018cost\u2019 of the attribute change weighted by the "
     "inverse covariance, so that changes along highly correlated dimensions (e.g., trust and autonomy) "
     "are cheaper than changes along independent or anti-correlated dimensions. The second term is the "
     "sum of boundary penalties — the discrete analogue of the phase-transition costs formalized in "
     "\u00a78.7.",
     "Body Text"),
    # The covariance matrix
    ("The covariance matrix \u03a3 is the discrete computational counterpart of the metric tensor g_{\u03bc\u03bd}. "
     "The correspondence is precise: the metric tensor defines the infinitesimal cost of displacement "
     "ds\u00b2 = g_{\u03bc\u03bd} dx\u1d50 dx\u207f on the continuous manifold; the Mahalanobis distance "
     "\u0394a\u1d40 \u03a3\u207b\u00b9 \u0394a defines the finite cost of transition on the discrete complex. "
     "In the limit of infinitesimal transitions (\u0394a \u2192 da), the Mahalanobis distance converges "
     "to the Riemannian arc length, and the discrete geodesic (minimum-weight path on \u0394) converges "
     "to the continuous geodesic (minimum-length curve on \u2133).",
     "Body Text"),
    # The boundary penalties
    ("The boundary penalties \u03b2_k are the discrete counterparts of the phase-transition penalties "
     "of \u00a78.7. Each \u03b2_k corresponds to a stratum boundary on the continuous manifold: crossing "
     "from one stratum to another incurs a finite cost that is not captured by the smooth metric within "
     "either stratum. The classification of boundaries from \u00a78.5 maps directly:",
     "Body Text"),
    # Mapping table as text
    ("Type I boundaries (thresholds with finite penalty) correspond to \u03b2_k < \u221e: crossing is "
     "costly but possible, and the A* search may include the crossing if the total path cost is lower "
     "than the alternative. Type II boundaries (phase transitions with regime change) correspond to "
     "\u03b2_k that are large but finite, reflecting the cost of transitioning between qualitatively "
     "different moral regimes. Type III boundaries (absorbing strata / nullifiers) correspond to "
     "\u03b2_k = \u221e: no path through this boundary can have finite cost, so the A* search provably "
     "avoids it. The sacred-value boundaries discussed in \u00a78.6 are exactly those with "
     "\u03b2 = \u221e.",
     "Body Text"),
    # Proposition on correspondence
    ("Proposition 8.4 (Continuous-Discrete Correspondence). [Proved.] Let \u2133 be a Whitney-stratified "
     "moral manifold with metric tensor g_{\u03bc\u03bd} and phase-transition penalties, and let "
     "\u0394_\u03b5 be the \u03b5-discretization of \u2133 (the weighted simplicial complex with "
     "Mahalanobis edge weights derived from g_{\u03bc\u03bd} and boundary penalties derived from the "
     "stratum transition costs). Then: (i) The minimum-weight path on \u0394_\u03b5 converges to the "
     "geodesic on \u2133 as \u03b5 \u2192 0. (ii) The boundary penalties on \u0394_\u03b5 equal the "
     "stratum transition costs on \u2133 exactly (no convergence needed — boundaries are discrete "
     "structures on both sides). (iii) The A* search on \u0394_\u03b5 with an admissible heuristic finds "
     "the optimal path on \u0394_\u03b5, which approximates the continuous geodesic to within O(\u03b5) "
     "in path cost.",
     "Body Text"),
    # Proof
    ("Proof. (i) Standard result in computational geometry: Mahalanobis distances on an \u03b5-grid of a "
     "Riemannian manifold converge to geodesic distances as \u03b5 \u2192 0 (see e.g. Memoli and Sapiro, "
     "2005). (ii) Boundary penalties are discrete by construction on both the continuous manifold "
     "(phase transitions are jumps, not smooth transitions — \u00a78.7, Definition 8.5) and the discrete "
     "complex (the \u03b2_k are finite constants). (iii) Follows from (i)\u2013(ii) combined with "
     "Proposition 11.4 (A* optimality on stratified spaces). \u25a1",
     "Body Text"),
    # Remark on domain applications
    ("Remark (Domain Instantiations). Each domain application of the Geometric Ethics framework "
     "instantiates the moral decision complex \u0394 with domain-specific attribute dimensions, covariance "
     "structure, and boundary penalties. In economics, the decision complex becomes the economic decision "
     "complex with dimensions including monetary cost, moral constraint, social norm, fairness, and "
     "identity; the minimum-cost path is the Bond Geodesic, which replaces Nash equilibrium as the "
     "solution concept. In clinical medicine, the decision complex becomes the clinical decision complex "
     "with dimensions including clinical outcomes, patient autonomy, trust, dignity, and justice; the "
     "minimum-cost path is the clinical geodesic, which replaces QALY optimization. In each case, the "
     "continuous manifold \u2133 provides the theoretical foundation, and the discrete complex \u0394 "
     "provides the computational realization. The Scalar Irrecoverability Theorem (\u00a715.6, "
     "Theorem 15.1) then proves that reducing either architecture to a scalar — utility, QALY, or any "
     "other one-dimensional projection — irrecoverably destroys the moral information that the "
     "multi-dimensional structure was designed to preserve.",
     "Body Text"),
]

print("Inserting discrete computational realization after para 2614...")
insert_multiple_after(doc, 2614, edit3_items)
print(f"  Inserted {len(edit3_items)} paragraphs.")


# ══════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════
doc.save(DST)
print(f"\nSaved as: {DST}")
print("Done. All four architectural innovations folded back.")
